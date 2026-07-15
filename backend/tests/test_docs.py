"""Тесты генерации документов (docgen), редактора docx (doc_editor), извлечения (extract)."""
import io
import re

from services import docgen
from services.extract import detect_kind, extract_text
from utils.doc_editor import apply_doc_edits, read_doc

SAMPLE = "# Договор\n**Жирный** абзац про XCMG.\n\n| A | B |\n|---|---|\n| 1 | 2 |"


def test_create_docx_is_valid():
    from docx import Document
    data = docgen.create_docx(SAMPLE, "test")
    assert len(data) > 1000
    doc = Document(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Договор" in text and "XCMG" in text


def test_create_pdf_is_valid():
    data = docgen.create_pdf("# Заголовок\nКириллический текст.", "test")
    assert data[:5] == b"%PDF-" and len(data) > 800


def test_doc_editor_read_and_edit_roundtrip():
    dx = docgen.create_docx("Первый абзац.\n\nВторой абзац.", "d")
    listing = read_doc(dx)
    assert "P1#" in listing or "параграф" in listing.lower()
    ref = re.search(r"P\d+#\w+", listing).group(0)
    out, diff = apply_doc_edits(dx, [{"op": "rewrite", "ref": ref, "new_text": "Изменённый абзац."}])
    assert len(out) > 1000
    from docx import Document
    text = "\n".join(p.text for p in Document(io.BytesIO(out)).paragraphs)
    assert "Изменённый абзац." in text


def test_create_proposal():
    from docx import Document
    data = {
        "filename": "kp", "title": "Поставка спецтехники", "client": "ООО «Стройка»",
        "markup_percent": 12, "validity_days": 14, "contact": "Иван · +7 900 000",
        "items": [
            {"name": "Экскаватор XCMG XE215C", "qty": 1, "price": 9850000},
            {"name": "Ковш дополнительный", "qty": 2, "price": 150000},
        ],
    }
    out = docgen.create_proposal(data)
    assert len(out) > 1000
    text = "\n".join(p.text for p in Document(io.BytesIO(out)).paragraphs)
    tables_text = " ".join(
        c.text for t in Document(io.BytesIO(out)).tables for row in t.rows for c in row.cells)
    assert "КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ" in tables_text
    assert "Стройка" in text
    assert "Экскаватор XCMG XE215C" in tables_text
    # экскаватор с наценкой 12%: 9 850 000 * 1.12 = 11 032 000
    assert "11 032 000" in tables_text
    # итог: 11 032 000 + 2*150 000*1.12 = 11 368 000
    assert "11 368 000" in tables_text
    # реквизиты компании присутствуют (issue #26)
    assert "2447007401" in text  # ИНН «Ремтехники»


def test_create_proposal_pdf():
    # issue #28 — КП в PDF
    data = {
        "filename": "kp", "title": "Поставка спецтехники", "client": "ООО «Стройка»",
        "markup_percent": 10, "items": [{"name": "Экскаватор", "qty": 1, "price": 1000000}],
    }
    out = docgen.create_proposal_pdf(data)
    assert out[:5] == b"%PDF-" and len(out) > 1000


def test_create_spec_report():
    # issue #25 — отчёт анализа ТЗ
    from docx import Document
    data = {
        "title": "ТЗ на портал заявок", "summary": "Веб-портал для приёма заявок клиентов.",
        "requirements": ["Авторизация сотрудников", "Личный кабинет клиента"],
        "risks": ["Не указаны сроки"], "contradictions": ["П.3 противоречит П.7"],
        "gaps": ["Не указан объём нагрузки"],
    }
    out = docgen.create_spec_report(data)
    d = Document(io.BytesIO(out))
    text = "\n".join(p.text for p in d.paragraphs)
    tables = " ".join(c.text for t in d.tables for r in t.rows for c in r.cells)
    assert "АНАЛИЗ" in tables and "Требования" in tables
    assert "Авторизация сотрудников" in text and "Не указаны сроки" in text


def test_create_estimate():
    # issue #27 — Excel-смета с настоящими формулами
    from openpyxl import load_workbook
    data = {
        "title": "Смета на ТО", "client": "ООО «Тест»", "markup_percent": 10,
        "items": [
            {"name": "Работа механика", "unit": "ч", "qty": 8, "price": 3600},
            {"name": "Масло моторное", "unit": "л", "qty": 20, "price": 450},
        ],
    }
    out = docgen.create_estimate(data)
    ws = load_workbook(io.BytesIO(out)).active   # формулы как строки
    cells = [str(c.value) for row in ws.iter_rows() for c in row if c.value is not None]
    joined = " ".join(cells)
    assert "Работа механика" in joined and "ИТОГО" in joined
    assert any(v.startswith("=ROUND") for v in cells)   # сумма позиции — формула
    assert any(v.startswith("=SUM(") for v in cells)     # итог — формула


def test_fill_template():
    # issue #26 — подстановка {{ПОЛЕ}} с сохранением структуры
    from docx import Document
    doc = Document()
    doc.add_paragraph("Договор с {{КЛИЕНТ}} на сумму {{ЦЕНА}} рублей.")
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Дата: {{ДАТА}}"
    t.rows[0].cells[1].text = "Поле {{НЕ_ЗАПОЛНЕНО}}"
    buf = io.BytesIO()
    doc.save(buf)

    out, filled, remaining = docgen.fill_template(
        buf.getvalue(), {"КЛИЕНТ": "ООО Ромашка", "ЦЕНА": "1 000 000", "ДАТА": "08.07.2026"})
    d = Document(io.BytesIO(out))
    text = "\n".join(p.text for p in d.paragraphs)
    tables = " ".join(c.text for tb in d.tables for r in tb.rows for c in r.cells)
    assert "ООО Ромашка" in text and "1 000 000" in text and "{{КЛИЕНТ}}" not in text
    assert "08.07.2026" in tables
    assert set(filled) == {"КЛИЕНТ", "ЦЕНА", "ДАТА"}
    assert remaining == ["НЕ_ЗАПОЛНЕНО"]


def _tiny_png(path):
    """Валидный маленький PNG (для проверки вставки логотипа, issue #26)."""
    from PIL import Image
    Image.new("RGB", (16, 16), (255, 203, 5)).save(str(path), "PNG")
    return str(path)


def test_logo_file_override_default_and_svg(tmp_path, monkeypatch):
    # issue #26 — LOGO_PATH переопределяет дефолт; SVG отклоняется; пусто → bundled
    import os

    from app.config import get_settings
    from services.docx_style import BUNDLED_LOGO, logo_file
    s = get_settings()
    png = _tiny_png(tmp_path / "logo.png")
    monkeypatch.setattr(s, "logo_path", png)
    assert logo_file() == png                                       # override PNG
    monkeypatch.setattr(s, "logo_path", str(tmp_path / "logo.svg"))  # SVG не годится
    assert logo_file() == ""
    monkeypatch.setattr(s, "logo_path", "")                          # пусто → поставляемый
    assert logo_file() == (BUNDLED_LOGO if os.path.isfile(BUNDLED_LOGO) else "")


def test_bundled_logo_exists_and_embeds():
    # issue #26 — поставляемый логотип на месте и реально встраивается в КП по умолчанию
    import os

    from docx import Document

    from services.docx_style import BUNDLED_LOGO
    assert os.path.isfile(BUNDLED_LOGO), "assets/logo.png отсутствует"
    out = docgen.create_proposal({"filename": "kp", "items": [{"name": "X", "price": 100}]})
    assert len(Document(io.BytesIO(out)).inline_shapes) >= 1


def test_proposal_embeds_logo_when_configured(tmp_path, monkeypatch):
    # issue #26 — при заданном LOGO_PATH логотип встраивается в КП картинкой
    from docx import Document

    import services.docx_style as ds
    png = _tiny_png(tmp_path / "logo.png")
    monkeypatch.setattr(ds, "logo_file", lambda: png)
    out = docgen.create_proposal({"filename": "kp", "items": [{"name": "X", "price": 100}]})
    doc = Document(io.BytesIO(out))
    assert len(doc.inline_shapes) >= 1   # логотип присутствует как встроенная картинка


def test_proposal_no_logo_without_config(monkeypatch):
    # без LOGO_PATH — КП генерируется как раньше, без картинок (мягкий пропуск)
    from docx import Document

    import services.docx_style as ds
    monkeypatch.setattr(ds, "logo_file", lambda: "")
    out = docgen.create_proposal({"filename": "kp", "items": [{"name": "X", "price": 100}]})
    assert len(Document(io.BytesIO(out)).inline_shapes) == 0


def test_detect_kind():
    assert detect_kind("a.docx") == "docx"
    assert detect_kind("b.PDF") == "pdf"
    assert detect_kind("c.png") == "image"
    assert detect_kind("d.xlsx") == "xlsx"
    assert detect_kind("e.unknown") == "other"


def test_extract_text_from_docx():
    dx = docgen.create_docx("Прайс на запчасти XCMG.", "p")
    text = extract_text(dx, "p.docx")
    assert "запчасти" in text
